import json
import ast
import os
import time
from unittest import result
from typing import Optional
import requests
import base64
import msal
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
import io
import urllib3

# Disable SSL warnings (for development only)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

from autogen_agentchat.agents import AssistantAgent, UserProxyAgent
from autogen_agentchat.teams import RoundRobinGroupChat, SelectorGroupChat
from autogen_agentchat.conditions import TextMentionTermination, MaxMessageTermination, SourceMatchTermination, TextMessageTermination
from autogen_ext.models.openai import AzureOpenAIChatCompletionClient
from autogen_agentchat.ui import Console
from autogen_core import CancellationToken
from autogen_agentchat.messages import TextMessage
from autogen_agentchat.base import TaskResult
from autogen_core.tools import FunctionTool
import asyncio


# Load environment variables
load_dotenv()

# CONFIGURATIONS
AZURE_CONFIG = {
    "tenant_id": os.getenv("AZURE_TENANT_ID"),
    "client_id": os.getenv("AZURE_CLIENT_ID"),
    "client_secret": os.getenv("AZURE_CLIENT_SECRET"),
    "sharepoint_domain": os.getenv("SHAREPOINT_DOMAIN"),
    "site_name": os.getenv("SHAREPOINT_SITE_NAME"),
    "site_id": os.getenv("SHAREPOINT_SITE_ID"),
    "folder_path": os.getenv("SHAREPOINT_FOLDER_PATH"),
    "drive_name": os.getenv("SHAREPOINT_DRIVE_NAME"),
    "sender_email":os.getenv("SENDER_EMAIL"),
    "support_email":os.getenv("SUPPORT_EMAIL"),
    "drive_id":os.getenv("DRIVE_ID"),
    "review_folder_url": os.getenv("REVIEW_FOLDER_URL"),
    "approved_folder_url": os.getenv("APPROVED_FOLDER_URL")
}

# Azure OpenAI model client
oai = AzureOpenAIChatCompletionClient(
    azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
    model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
)

# Config: Salesforce Connected App details from environment variables
SF_CLIENT_ID = os.getenv("SF_CLIENT_ID")
SF_CLIENT_SECRET = os.getenv("SF_CLIENT_SECRET")
SF_LOGIN_URL = os.getenv("SF_LOGIN_URL")
SF_INSTANCE_URL = os.getenv("SF_INSTANCE_URL")
SF_API_VERSION = os.getenv("SF_API_VERSION")

# Validate that all required environment variables are set
required_vars = {
    "SF_CLIENT_ID": SF_CLIENT_ID,
    "SF_CLIENT_SECRET": SF_CLIENT_SECRET,
    "SF_LOGIN_URL": SF_LOGIN_URL,
    "SF_INSTANCE_URL": SF_INSTANCE_URL,
    "SF_API_VERSION": SF_API_VERSION
}

missing_vars = [var for var, value in required_vars.items() if not value]
if missing_vars:
    raise ValueError(f"Missing required environment variables: {', '.join(missing_vars)}")

# ================== FUNCTIONS ==================

def msal_token() -> str:
    client_id = AZURE_CONFIG["client_id"]
    tenant_id = AZURE_CONFIG["tenant_id"]
    client_credentials = AZURE_CONFIG["client_secret"]
    authority = f"https://login.microsoftonline.com/{tenant_id}"

    app = msal.ConfidentialClientApplication(
        client_id,
        authority=authority,
        client_credential=client_credentials,
        verify=False  # Disable SSL verification
    )

    scope = ["https://graph.microsoft.com/.default"]
    token_response = app.acquire_token_for_client(scopes=scope)

    if "access_token" in token_response:
        return token_response["access_token"]
    else:
        print("Failed to get token:", token_response.get("error_description"))

def get_list_items(access_token: str, site_id: str, list_name: str = "Test_New_Salesforce Accounts Data", column_name: str = "field_1") -> list:
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/lists/{list_name}/items?expand=fields"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers, verify=False)
    if response.status_code == 200:
        items = response.json().get("value", [])
        customer_list = [item.get("fields", {}).get(column_name) for item in items]
        return customer_list
    else:
        raise Exception(f"❌ Error fetching list items: {response.status_code} {response.text}")

def get_customer_list():
    access_token = msal_token()
    site_id = AZURE_CONFIG["site_id"]
    try:
        customer_list = get_list_items(access_token, site_id)
        return {
            "status": "success",
            "customer_list": customer_list,
            "count": len(customer_list),
            "message": "Customer list retrieved successfully. Please compare the extracted name from the summary with this list."
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to retrieve customer list: {str(e)}"}

def load_transcript_from_sharepoint(file_name: str):
    access_token = msal_token()
    site_id = AZURE_CONFIG["site_id"]
    drive_id = "b!So-TjS5i9Uewx7vi-XtNSf_SkM1pp49In_9GWcS1ynufpl5vClqxT506ePyz_Qyc"
    folder_name = "Initial Transcripts"
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{folder_name}/{file_name}:/content"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers, verify=False)
    if response.status_code == 200:
       return {"transcript": response.text, "File name": file_name}
    else:
        print(f"Error: {response.status_code} - {response.text}")
        return "Error"
    

# ================== SITE MEMBERSHIP CHECK FUNCTION ==================

#IS USER MEMBER OF SITE

def is_user_member_of_site(user_email: str, site_name: str = "Test") -> dict:
    """
    Check if a user is a member of a SharePoint site by checking their group memberships.
    This works for both classic and modern SharePoint sites.
    
    The function checks if the user belongs to any Azure AD group that matches the site name.
    SharePoint sites typically create corresponding Azure AD groups with the same name.
    
    Args:
        user_email: Email address of the user to check
        site_name: Name of the SharePoint site (default: "Test")
    
    Returns:
        Dictionary with membership status and details:
        {
            "status": "success" | "error",
            "is_member": True | False,
            "site_name": str,
            "user_email": str,
            "user_display_name": str,
            "matching_group": {"name": str, "id": str} or None,
            "all_groups": [str],
            "message": str
        }
    """
    graph_token = msal_token()
    
    if not graph_token:
        return {
            "status": "error",
            "message": "Failed to acquire Graph API token",
            "is_member": False,
            "site_name": site_name,
            "user_email": user_email
        }

    headers = {
        "Authorization": f"Bearer {graph_token}",
        "Accept": "application/json"
    }

    print("=" * 70)
    print(f"🔍 Checking Site Membership")
    print(f"👤 User: {user_email}")
    print(f"📍 Site: {site_name}")
    print("=" * 70)

    try:
        # Step 1: Get user ID
        print(f"\n📌 Getting user information...")
        user_url = f"https://graph.microsoft.com/v1.0/users/{user_email}"
        user_response = requests.get(user_url, headers=headers, verify=False)
        
        if user_response.status_code != 200:
            print(f"   ❌ User not found")
            return {
                "status": "error",
                "message": f"User '{user_email}' not found",
                "is_member": False,
                "site_name": site_name,
                "user_email": user_email
            }
        
        user_data = user_response.json()
        user_id = user_data.get("id")
        display_name = user_data.get("displayName")
        print(f"   ✅ User: {display_name}")
        print(f"      ID: {user_id}")
        
        # Step 2: Get all groups the user belongs to
        print(f"\n📌 Checking user's group memberships...")
        member_of_url = f"https://graph.microsoft.com/v1.0/users/{user_id}/memberOf"
        member_of_response = requests.get(member_of_url, headers=headers, verify=False)
        
        if member_of_response.status_code != 200:
            print(f"   ❌ Failed to get group memberships")
            return {
                "status": "error",
                "message": "Failed to retrieve group memberships",
                "is_member": False,
                "site_name": site_name,
                "user_email": user_email,
                "user_display_name": display_name
            }
        
        groups = member_of_response.json().get("value", [])
        print(f"   Found {len(groups)} group(s)")
        
        # Step 3: Check if any group matches the site name
        is_member = False
        matching_group = None
        all_groups = []
        
        for group in groups:
            group_display_name = group.get("displayName")
            group_id = group.get("id")
            
            # Only add to list if displayName is not None
            if group_display_name:
                all_groups.append(group_display_name)
            
            # Check if group name matches site name (case-insensitive)
            if group_display_name and group_display_name.lower() == site_name.lower():
                is_member = True
                matching_group = {
                    "name": group_display_name,
                    "id": group_id
                }
                print(f"   ✅ MATCH FOUND: User is in group '{group_display_name}'")
        
        # Final result
        print("\n" + "=" * 70)
        if is_member:
            print(f"✅ YES - {user_email} IS A MEMBER of site '{site_name}'")
            print(f"   Via group: {matching_group['name']}")
        else:
            print(f"❌ NO - {user_email} is NOT a member of site '{site_name}'")
            print(f"   User's groups: {', '.join(all_groups)}")
        print("=" * 70)
        
        return {
            "status": "success",
            "is_member": is_member,
            "site_name": site_name,
            "user_email": user_email,
            "user_display_name": display_name,
            "matching_group": matching_group,
            "all_groups": all_groups,
            "message": f"User {'IS' if is_member else 'is NOT'} a member of site '{site_name}'"
        }
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": f"Error: {str(e)}",
            "is_member": False,
            "site_name": site_name,
            "user_email": user_email
        }

#CHECK SITE MEMBERSHIP AND GET FOLDER

def check_site_membership_and_get_folder(user_email: str, site_name: str = None) -> bool:
    """
    Check if a user is a member of the SharePoint site.
    
    Args:
        user_email: Email address of the user to check
        site_name: Name of the SharePoint site (defaults to SHAREPOINT_SITE_NAME from env)
    
    Returns:
        True if user is a member, False otherwise
    """
    # Use site name from environment if not provided
    if site_name is None:
        site_name = AZURE_CONFIG["site_name"]
    try:
        print(f"\n🔍 Checking site membership for user: {user_email}")
        
        # Call the GraphAPI function to check membership
        result = is_user_member_of_site(user_email, site_name)
        
        # Return membership status
        is_member = result.get("status") == "success" and result.get("is_member")
        
        if is_member:
            print(f"✅ User {user_email} is a member of {site_name}")
        else:
            print(f"❌ User {user_email} is NOT a member of {site_name}")
        
        return is_member
    
    except Exception as e:
        print(f"❌ Error checking site membership: {str(e)}")
        # Default to non-member in case of errors
        return False

# ================== FORMAT & UPLOAD FUNCTIONS ==================

#In Progress
#Approved Transcripts
def format_summary(summary: str, metadata: dict, sub_folder: str, user_email: Optional[str] = None):
    try:
        doc = Document()
        customer_name = metadata.get('Name of customer', 'Unknown Customer')
        doc.add_heading(f"Name of customer - {customer_name}", 0)
        meeting_date = metadata.get('Date of Meeting', 'Unknown Date')
        doc.add_paragraph(f"Date of Meeting: {meeting_date}")
        contact_info = metadata.get('Name & Designation of the customer whom you meet', 'Unknown Contact')
        doc.add_paragraph(f"Name & Designation: {contact_info}")
        doc.add_paragraph("Next Steps:")
        next_steps = metadata.get('Next Steps', 'No next steps defined')
        doc.add_paragraph(next_steps)
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)
        output_file = f"meeting_summary_{customer_name.replace(' ', '_')}.docx"
        status = upload_to_sharepoint(output_file, doc, sub_folder, user_email=user_email)
        if status["status"] == "Failed":
            return "UPLOADING FAILED: Successfully formatted the document but failed to upload"
        elif status["status"] == "Uploaded Successfully":
            return {'file_name':output_file, 'account_name': customer_name, 'folder_path': sub_folder}
    except Exception as e:
        print(f"❌ Error in formatting: {str(e)}")
        return {"error": f"Formatting failed: {str(e)}"}

def upload_to_sharepoint(file_name: str, doc, sub_folder:str, sharepoint_details: dict = None, user_email: Optional[str] = None):
    try:
        if sharepoint_details is None:
            sharepoint_details = AZURE_CONFIG
        required_keys = ["tenant_id", "client_id", "client_secret"]
        missing_keys = [key for key in required_keys if key not in sharepoint_details or sharepoint_details[key] == f"your-{key.replace('_', '-')}-here"]
        if missing_keys:
            return {
                "status": "Failed",
                "error": f"Please update AZURE_CONFIG with your actual: {', '.join(missing_keys)}"
            }
        authority = f"https://login.microsoftonline.com/{sharepoint_details['tenant_id']}"
        app = msal.ConfidentialClientApplication(
            sharepoint_details['client_id'],
            authority=authority,
            client_credential=sharepoint_details['client_secret'],
            verify=False  # Disable SSL verification
        )
        scopes = ["https://graph.microsoft.com/.default"]
        token_response = app.acquire_token_for_client(scopes=scopes)
        if "access_token" not in token_response:
            return {"status": "Failed", "error": "Failed to acquire access token", "details": token_response.get("error_description", "")}
        access_token = token_response["access_token"]
        site_id = AZURE_CONFIG["site_id"]
        drive_id = "b!So-TjS5i9Uewx7vi-XtNSf_SkM1pp49In_9GWcS1ynufpl5vClqxT506ePyz_Qyc"
        #sub_folder = "Approved Transcripts"
        upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{sub_folder}/{file_name}:/content"
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/octet-stream"
        }
        file_stream = BytesIO()
        doc.save(file_stream)
        file_content = file_stream.getvalue()
        response = requests.put(upload_url, headers=headers, data=file_content, verify=False)
        if response.status_code in [200, 201]:
            response_data = response.json()
            file_id = response_data.get("id", "")
            
            # Update UserEmail column if user_email is provided
            if user_email and file_id:
                update_sharepoint_item_user_email(drive_id, file_id, user_email, access_token)
            
            return {
                "status": "Uploaded Successfully",
                "file_name": file_name,
                "sharepoint_url": response_data.get("webUrl", ""),
                "file_id": file_id
            }
        else:
            return {
                "status": "Failed",
                "error": f"Upload failed with status {response.status_code}",
                "details": response.text
            }
    except Exception as e:
        error_msg = str(e)
        print(f"❌ SharePoint upload failed: {error_msg}")
        return {"status": "Failed", "file_path": file_name, "error": error_msg}


# Function to update UserEmail column in SharePoint after document upload
def update_sharepoint_item_user_email(drive_id: str, item_id: str, user_email: str, access_token: str) -> dict:
    """
    Updates the UserEmail column of a SharePoint item after upload.
    
    Args:
        drive_id: The document library's drive ID
        item_id: The DriveItem ID of the uploaded file
        user_email: The email of the user who uploaded the document
        access_token: Graph API access token
    
    Returns:
        Dictionary with status and updated fields or error message
    """
    try:
        url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/listItem/fields"
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        payload = {
            "UserEmail": user_email
        }
        response = requests.patch(url, headers=headers, json=payload, timeout=30, verify=False)
        if response.status_code == 200:
            print(f"✅ UserEmail column updated successfully for item {item_id}")
            return {"status": "success", "updated_fields": response.json()}
        else:
            print(f"❌ Failed to update UserEmail column: {response.status_code} {response.text}")
            return {"status": "failed", "error": f"Update failed: {response.status_code} {response.text}"}
    except Exception as e:
        print(f"❌ Error updating UserEmail column: {str(e)}")
        return {"status": "failed", "error": str(e)}


#Function to read document from sharepoint
def load_document_from_sharepoint(file_name: str, account_name:str):
    access_token = msal_token()
    site_id = AZURE_CONFIG["site_id"]
    drive_id = AZURE_CONFIG["drive_id"]
    folder_name = "Approved Transcripts"
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{folder_name}/{file_name}:/content"
    headers = {"Authorization": f"Bearer {access_token}"}
    response = requests.get(url, headers=headers, verify=False)
    if response.status_code == 200:
        file_bytes = BytesIO(response.content)
        doc = Document(file_bytes)
        text_content = "\n".join([para.text for para in doc.paragraphs])
       # print("\n\nWord document content - \n",text_content,"\n")
        file_bytes.seek(0)
        # return {f"File Name": {file_name},
        #         "File Bytes":{file_bytes}}
        print("\n\nLoaded the file, now passing it to upload to salesforce")
        return upload_to_salesforce(file_bytes, account_name, file_name, text_content)
    else:
        print(f"Error: {response.status_code} - {response.text}")
        return "TERMINATE: DOCUMENT FAILED TO LOAD"
    
#Function to upload the document to salesforce
def upload_to_salesforce(doc : bytes, account_name:str, file_name:str, word_doc_content:str):

        """
        Attaches a file to an Account in Salesforce using ContentVersion API (Client Credentials flow).
        """

        if doc != None:
           #doc_content = base64.b64encode(doc.read()).decode("utf-8")
           doc_content = word_doc_content
 

        # Step 1: Get access token (Client Credentials Flow)
        token_url = f"{SF_LOGIN_URL}/services/oauth2/token"
        payload = {
            "grant_type": "client_credentials",
            "client_id": SF_CLIENT_ID,
            "client_secret": SF_CLIENT_SECRET
        }
        resp = requests.post(token_url, data=payload, verify=False)
        resp.raise_for_status()
        access_token = resp.json()["access_token"]

        # Step 2: Query Account Id by Name
        query = f"SELECT Id FROM Account WHERE Name = '{account_name}' LIMIT 1"
        query_url = f"{SF_INSTANCE_URL}/services/data/v{SF_API_VERSION}/query"
        headers = {"Authorization": f"Bearer {access_token}"}
        resp = requests.get(query_url, headers=headers, params={"q": query}, verify=False)
        resp.raise_for_status()
        records = resp.json().get("records", [])
        if not records:
            raise Exception(f"Account with name '{account_name}' not found.")
        account_id = records[0]["Id"]

        # Step 3: Upload file using ContentVersion
        upload_url = f"{SF_INSTANCE_URL}/services/data/v{SF_API_VERSION}/sobjects/ContentVersion"
        #read the file from the sharepoint approved folder path

        entity_content = {
            "Title": file_name,
            "PathOnClient": file_name,
            "FirstPublishLocationId": account_id
        }

        # Open file and ensure it gets closed properly
    
        files = {
                "entity_content": (None, json.dumps(entity_content), "application/json"),
                "VersionData": (file_name, doc, "application/octet-stream")
            }

       
        resp = requests.post(upload_url, headers={"Authorization": f"Bearer {access_token}"}, files=files, verify=False)
        resp.raise_for_status()
      

        result = resp.json()
        
        if result.get("success") and result.get("id"):
            return {"message":"SUCCESSFULLY UPLOADED THE DOCUMENT IN SALESFORCE",
                    "Document content": doc_content,
                     "File name": file_name}
        else:
              print("ERROR : No Document Found")
              return "TERMINATE: NO DOCUMENT FOUND"


        

    # if doc != None:
    #    docx_base64 = base64.b64encode(doc.read()).decode("utf-8")
    #    print("\n\nTEXT CONTENT = \n\n",docx_base64)
    #    return "SUCCESSFULLY UPLOADED THE DOCUMENT IN SALESFORCE"
    # else:
    #     print("ERROR : No Document Found")
    #     return "TERMINATE: NO DOCUMENT FOUND"



#Function to call power automate flow
def call_power_automate(flow_url : str = "https://default7eb95c17398a40349e24561c1f9a12.eb.environment.api.powerplatform.com:443/powerautomate/automations/direct/workflows/4c3dbba6cdef4a0e8bc28c6c49781aea/triggers/manual/paths/invoke?api-version=1&sp=%2Ftriggers%2Fmanual%2Frun&sv=1.0&sig=wMUais-ZKWv0KxSX5LQhrFJXK_nzjL02-Jyme_gWMH0",
                        payload: dict = {
    "TranscriptSummaryText": "1. Date of the meeting: August 15, 2025 \n 2. Name of customer: NA \n 3. Name & Designation of the customer whom you meet: Not mentioned in the transcript \n4. What are the key priorities for the organization: Upcoming project on data and ML activities\n5. What are they key priorities of the individual: Not explicitly mentioned\n6. What are areas of interest where Apexon can contribute: Providing a POC on ADF (Azure Data Factory)\n7. What are the next steps: Schedule the next meeting on August 25,2025",
    "UserEmail": AZURE_CONFIG["sender_email"]
}):
    """
    Triggers a Power Automate flow via its HTTP URL.
    :param flow_url: The Power Automate trigger URL
    :param payload: Data you want to send to the flow (JSON)
    """
    try:
        headers = {"Content-Type": "application/json"}
        response = requests.post(flow_url, headers=headers, data=json.dumps(payload), verify=False)
        
        # Optional: If the flow doesn't expect data, you can just do:
        # response = requests.post(flow_url)
        
        if response.status_code in (200, 202):
            print("✅ Power Automate flow triggered successfully!")
            #print("Response:", response.text)
        else:
            print (f"❌ Failed to trigger flow. Status Code: {response.status_code}")
            #print("Response:", response.text)
    except Exception as e:
        return f"❌ Error while calling Power Automate:, {str(e)}"


#Function to send a email
def send_email(text: Optional[str] = None, file_name: Optional[str] = None, email_body: Optional[str] = None, to: Optional[str] = None, subject: Optional[str] = None):
    token = msal_token()

    if text is not None:

        if ".txt" in file_name:
            # ✅ Convert plain text into bytes
            text_bytes = text.encode("utf-8")
            text_base64 = base64.b64encode(text_bytes).decode("utf-8")
            content_type = "Text"
        elif ".docx" in file_name:
            doc = Document()
            doc.add_paragraph(text)  # your plain text

            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
        
            text_base64 = base64.b64encode(file_stream.read()).decode("utf-8")
            content_type = "HTML"


        email_msg = {
            "message": {
                "subject":subject,
                "body":{
                    "contentType":content_type,
                    "content":email_body
                    },
                "toRecipients":  [
                {"emailAddress": {"address": to}}
            ],
            "attachments": [
                    {
                        "@odata.type": "#microsoft.graph.fileAttachment",
                        "name": file_name,
                        "contentBytes": text_base64
                    }
                ]
            },
            "saveToSentItems":"true"
        }
    else:
        email_msg = {
            "message": {
                "subject":subject,
                "body":{
                    "contentType":"HTML",
                    "content":email_body
                    },
                "toRecipients":  [
                {"emailAddress": {"address": to}}
            ],
            },
            "saveToSentItems":"true"
        }


    sender = AZURE_CONFIG["sender_email"]

    # API endpoint for sending mail
    #url = f"https://graph.microsoft.com/v1.0/users/Rohan.Patil@isplahdplayground.onmicrosoft.com/sendMail"
    url = "https://graph.microsoft.com/v1.0/users/Rohan.Patil@isplahdplayground.onmicrosoft.com/sendMail"

    

    # Send request
    response = requests.post(
    url,
    headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
    json=email_msg,
    verify=False)
    
    if response.status_code == 202:
        print("✅ Email sent successfully!")
        return subject;
    else:
        print("❌ Error:", response.status_code, response.text)
        return subject;



#Function which moves a file from one folder to another
def move_file_in_sharepoint(filename:str = "move.txt", destinationFolderName:str = "In Progress", sourceFolderName:str = "Approved Transcripts"):
    token = msal_token()

    
    # Step 1: Get the file ID from the source folder
    search_url = f"https://graph.microsoft.com/v1.0/sites/{AZURE_CONFIG['site_id']}/drives/{AZURE_CONFIG['drive_id']}/root:/{sourceFolderName}/{filename}"
    search_url += ":"

    search_response = requests.get(search_url, headers={"Authorization": f"Bearer {token}"}, verify=False)

    if search_response.status_code != 200:
        print(f"❌ File not found: {filename} in {sourceFolderName}")
        return

    file_info = search_response.json()
    item_id = file_info["id"]  # Unique ID of the file in SharePoint

    # Step 2: Move file by updating its parentReference
    move_url = f"https://graph.microsoft.com/v1.0/sites/{AZURE_CONFIG['site_id']}/drives/{AZURE_CONFIG['drive_id']}/items/{item_id}"
    move_body = {
        "parentReference": {
            "path": f"/drives/{AZURE_CONFIG['drive_id']}/root:/{destinationFolderName}"
        }
    }

    move_response = requests.patch(
        move_url,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        },
        json=move_body,
        verify=False
    )

    if move_response.status_code in (200, 201):
        print(f"✅ File '{filename}' moved from '{sourceFolderName}' → '{destinationFolderName}'")
    else:
        print(f"❌ Failed to move file. Status: {move_response.status_code}")
        print(move_response.text)


#Function to delete a file from sharepoint
def delete_file_from_sharepoint(source_folder_name:str = "Approved Transcripts", file_name:str = "meeting_summary_UrbanTech_Solutions.docx"):
    """
    Delete a file from a SharePoint folder using Microsoft Graph API.
    """

    access_token = msal_token()
    site_id = AZURE_CONFIG["site_id"]
    drive_id = AZURE_CONFIG["drive_id"]




    # Step 1: Build file path URL
    file_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/root:/{source_folder_name}/{file_name}"

    # Step 2: Get file ID first
    get_file_response = requests.get(
        file_url,
        headers={"Authorization": f"Bearer {access_token}"},
        verify=False
    )

    if get_file_response.status_code != 200:
        print (f"❌ File not found or folder incorrect. Details: {get_file_response.text}")

    file_id = get_file_response.json()['id']

    # Step 3: Delete the file
    delete_url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives/{drive_id}/items/{file_id}"

    delete_response = requests.delete(
        delete_url,
        headers={"Authorization": f"Bearer {access_token}"},
        verify=False
    )

    if delete_response.status_code == 204:
        print (f"✅ File '{file_name}' deleted successfully from folder '{source_folder_name}'.")
    else:
        print (f"❌ Failed to delete file. Details: {delete_response.text}")






# ================== WRAP FUNCTIONS INSIDE TOOLS ==================
   
get_customer_list_tool =  FunctionTool(get_customer_list, description="Gets a list of customer name from a sharepoint site")
check_site_membership_tool = FunctionTool(check_site_membership_and_get_folder, description="Checks if a user is a member of the SharePoint site and returns True if member, False if not")
format_summary_tool = FunctionTool(format_summary, description="Formats the summary inside a document and uploads the document to a sharepoint site. Accepts summary, metadata, sub_folder, and optional user_email to update UserEmail column in SharePoint.")
load_document_from_sharepoint_tool = FunctionTool(load_document_from_sharepoint, description="Loads a document(docx) file from a sharepoint site")
call_power_automate_tool = FunctionTool(call_power_automate, description="Triggers a power automate flow")
send_email_tool = FunctionTool(send_email, description="Sends a email.")
move_file_in_sharepoint_tool = FunctionTool(move_file_in_sharepoint, description = "Moves a file from one sharepoint folder to another.")
delete_file_from_sharepoint_tool = FunctionTool(delete_file_from_sharepoint,description="Deletes a particular file from a sharepoint folder.")

# ================== AGENTS ==================

user_proxy = UserProxyAgent("User")


summarizer_agent = AssistantAgent(
    "Summarizer",
    system_message="""
You are a meeting transcript summarizer. You will get a meeting transcript

Step 1: Always generate a summary from the transcript in the following format:

1. Date of the meeting
2. Name of customer
3. Name & Designation of the customer whom you meet
4. What are the key priorities for the organization
5. What are they key priorities of the individual
6. What are areas of interest where Apexon can contribute
7. What are the next steps

**IMPORTANT:**
- When extracting the customer name, use the exact text as it appears in the transcript.
- Do NOT correct, guess, or change the customer name in any way.
- If the name is unclear, misspelled, or abbreviated, copy it as it is.
- Do NOT use your own knowledge or make up a name.

Examples:
- If transcript says "TFWW", summary must say "2. Name of customer: TFWW"
- If transcript says "Billows", summary must say "2. Name of customer: Billows"
- If transcript says "Tforce World Wide", summary must say "2. Name of customer: Tforce World Wide"
- If transcript says "Tforse", summary must say "2. Name of customer: Tforse"
- If transcript says "Xyzabc", summary must say "2. Name of customer: Xyzabc"

Also include the original meeting transcript at the end.

Return ONLY the final summary text with original meeting transcript without any modification. Do not explain anything apart from the summary.
""",
    model_client=oai,
    description="Summerizes the given transcript text then continues the workflow by calling GetCustomerListAgent"
)

GetCustomerListAgent = AssistantAgent(
    "GetCustomerList",
    system_message="""You are a intelligent Get customer list assistant agent.
You have access to get_customer_list function which returns customer list from sharepoint.
You MUST call the get_customer_list function and get the customer list from the sharepoint.
After getting the customer list from the sharepoint, Return ONLY a valid Python dictionary with EXACTLY two keys:
- 'summary': the corrected summary text recieved from summarizer agent after summarizing the transcript (unaltered, just pass it through)
- 'Customer list': the whole customer list received from get_customer_list function.""",
    model_client=oai,
    tools=[get_customer_list_tool],
    description="Loads the list of customer names from a sharepoint site"
)

GetCorrectCustomerName_agent = AssistantAgent(
    "GetCorrectCustomerName",
    system_message="""You are an intelligent GetCorrectCustomerName assistant agent.

You are given a:

1. A **summary** generated from a voice recording.
2. A customer list - Received from GetCustomerListAgent.

Your task is to Extract the customer name from the summary, the customer name will follow this pattern:
  "2) Name of customer: <original_name>" .

  Find the most likely match of <original_name> from the customer list based on:
- Pronunciation similarity, including silent letters (e.g., Knipper sounds like Nipper)
- Common abbreviations (e.g., TFWW → Tforce World Wide)
- Phonetic errors or misspellings from speech-to-text
- Informal or partial references (e.g., Billows → Optimizely - Billows Electric Supply Co)

**Instructions:**
If a close match exists,  replace the value after "2) Name of customer:" in the summary with the full correct name from the list.
- Keep the rest of the summary and original meeting transcript completely unchanged.
- If no match is found and if you are **not confident**, return the **originalName** exactly as it was given (do not guess and change).
- Do NOT make up or hallucinate a name.

Examples:
• Input: TFWW → Output: Tforce World Wide
• Input: Billows → Output: Optimizely - Billows Electric Supply Co
• Input: Eastern → Output: Eastern Bank
• Input: Somerandomcompany → Output: Somerandomcompany   ← (No match found)

Return only the final summary text with original meeting transcript without any modification. Do not explain anything.""",
    model_client=oai,
    description="If the customer name is misspelled, it corrects the name by finding the closest match from the customer list."
)

extractor_agent = AssistantAgent(
    "Extractor",
    system_message="""You are an extractor agent.
Your job is to carefully read the given meeting summary and extract structured metadata.

Return ONLY a valid Python dictionary with EXACTLY two keys:
- 'summary': the corrected summary text recieved from GetCorrectCustomerName_agent after correcting the customer name  (unaltered, just pass it through)
- 'metadata': another dictionary with EXACTLY these keys:
    - 'Date of Meeting'
    - 'Name of customer'
    - 'Name & Designation of the customer whom you meet'
    - 'Key priorities for the organization'
    - 'Key priorities of the individual'
    - 'Areas of interest where Apexon can contribute'
    - 'Next Steps'

Make sure you also include original meeting transcript also after the summary.
STRICT and IMPORTANT INSTRUCTION - 1.)  You MUST return both the keys - summary as well as metadata to FolderPathDecisionAgent.
                                   2.) YOU MUST include original transcript also after summary.
Do not include anything else in your response. Do not use code blocks.""",
    model_client=oai,
    description="Extracts structured metadata and returns summary and metadata in python dictionary format"
)

FolderPathDecisionAgent = AssistantAgent(
    "FolderPathDecisionAgent",
    system_message="""You are a Site Membership Check agent. Your task is to check if the logged-in user is a member of the SharePoint site and determine the correct upload folder.

    STEP 1:
        You will receive a python dictionary with 'summary' and 'metadata' from the Extractor agent.
        You MUST preserve this dictionary as you will need to pass it to UploadToSharepointAgent.
    
    STEP 2:
        You MUST call the check_site_membership_and_get_folder function with the user_email parameter.
        The user email will be provided in the workflow instructions.
        This function will return True or False:
        - True: user is a member
        - False: user is NOT a member
    
    STEP 3:
        Based on the True/False result from the function, YOU must decide the folder_path:
        - If True (member): folder_path = "Approved Transcripts"
        - If False (non-member): folder_path = "Approved Transcripts In Review"
    
    STEP 4:
        After deciding the folder_path, you MUST pass FOUR things to UploadToSharepointAgent:
        1. summary (from Extractor)
        2. metadata (from Extractor)
        3. folder_path (that YOU decided based on membership result)
        4. user_email (the email used for membership check - provided in workflow instructions)
        
        Format your response to UploadToSharepointAgent as:
        "Here is the data for upload: summary: [summary_text], metadata: [metadata_dict], folder_path: [folder_path_value], user_email: [user_email_value]"
        
    Do not modify the summary or metadata. Check membership, decide the folder path, and pass all required data forward.""",
    model_client=oai,
    tools=[check_site_membership_tool],
    description="Checks user's site membership and determines appropriate upload folder"
)

UploadToSharepointAgent = AssistantAgent(
    "UploadToSharepointAgent",
    system_message="""You are a UploadToSharepointAgent agent. You have access to a format_summary function which has 4 input parameters: summary, metadata, sub_folder, and user_email.

    STEP 1:
        You will receive FOUR pieces of information from FolderPathDecisionAgent:
        1. summary (string) - the meeting summary
        2. metadata (python dictionary) - extracted metadata
        3. folder_path (string) - either "Approved Transcripts" or "Approved Transcripts In Review"
        4. user_email (string) - the email of the logged-in user
        
        Your task is to format the summary in a word document and upload it to the specified SharePoint folder.
        
        You MUST call the format_summary function with ALL FOUR arguments:
        - "summary" (string)
        - "metadata" (dictionary)
        - "sub_folder" (string) - use the folder_path received from FolderPathDecisionAgent
        - "user_email" (string) - use the user_email received from FolderPathDecisionAgent
        
        Do not call format_summary with only two or three arguments. All four parameters are required.
        The user_email will be used to update the UserEmail column in SharePoint after the document is uploaded.

    STEP 2:
        After executing and uploading the word document to sharepoint, you will receive a python dictionary response from the format_summary tool.
        This dictionary contains: file_name, account_name, and folder_path.
        Return that COMPLETE response (including folder_path) as it is to Salesforce_uploader agent.
        """,
    model_client=oai,
    tools=[format_summary_tool],
    description="Does formatting and uploads to sharepoint site in appropriate folder based on user membership, also updates UserEmail column"
)

salesforce_uploader_agent = AssistantAgent(
    "Salesforce_uploader",
     model_client=oai,
     tools=[load_document_from_sharepoint_tool],
     system_message=""" 

     "You are a salesforce uploader agent. You will receive a python dictionary response from UploadToSharepointAgent which contains 3 parameters - 1.)file_name 2.)account_name 3.)folder_path"
        
     IMPORTANT: First check the folder_path value received from UploadToSharepointAgent.
     
     IF folder_path contains "Review" (e.g., "Approved Transcripts In Review"):
        - DO NOT call the load_document_from_sharepoint function
        - DO NOT upload to Salesforce
        - Instead, immediately send this message to Email_agent:
          "Document uploaded to Review folder - Salesforce upload skipped. file_name: [file_name], folder_path: [folder_path]"
        - This is because non-members should not have their documents uploaded to Salesforce.
     
     ELSE (folder_path is "Approved Transcripts"):
        STEP 1: 
        "Your task is to read a docx file from a sharepoint site and upload the same file to salesforce."
        "You have access to load_document_from_sharepoint function which has 2 input parameters."
        "You MUST call the load_document_from_sharepoint function by passing the document(docx) file_name and account_name both received from UploadToSharepointAgent as input parameters."
        "ONLY use function to read and upload the docx file"
        
        STEP 2:
        "After uploading the document to salesforce, send this message to Email_agent:"
        "Document uploaded in the salesforce successfully. file_name: [file_name], folder_path: [folder_path]"
         CRITICAL - Make sure you call the load_document_from_sharepoint tool and upload it to salesforce if folder path is "Approved Transcripts" .


     If salesforce upload fails or if Account name is not found in the salesforce - then return "Account name is not found - TERMINATE" and stop the workflow strictly."
     """,


     description="Reads a document from sharepoint and uploads it to salesforce (only for site members)"
)


CallPowerAutomateAgent = AssistantAgent(
    "CallPowerAutomateAgent",
     model_client = oai,
     description = "Calls a power automate flow",
     system_message =""" "You are a Power automate agent, you have access to call_power_automate function." \
                      "Your task is to call the power automate workflow if the name of the customer is missing from the transcript."
                      "You MUST call the call_power_automate function to trigger the power automate flow" \
                      "Only use the function to call the power automate flow" \
                      "Return TERMINATE after calling the power automate workflow and stop this workflow" """,
     tools = [call_power_automate_tool]
)


# Email will be passed dynamically through workflow instructions

email_agent = AssistantAgent(
    "Email_agent",
    system_message=f""" "You are an email agent with access to send_email function.

                        You will receive a message from Salesforce_uploader agent containing file_name and folder_path.
                        
                        IMPORTANT: Check the folder_path value to determine which email to send:
                        
                        CASE 1 - If folder_path contains "Review" (non-member scenario):
                            Call send_email with these parameters:
                            1.) text = None (no attachment needed)
                            2.) file_name = None
                            3.) email_body = "Dear Support Team,<br><br>A new meeting summary document has been uploaded to the Review folder for your approval.<br><br>The document requires review as the user who submitted it is not a member of the {AZURE_CONFIG['site_name']} site.<br><br><a href='{AZURE_CONFIG['review_folder_url']}' target='_blank'>Link to review folder</a><br><br>Please review and approve the document.<br><br>Thanks,<br>ServiceAccount"
                            4.) to = "{AZURE_CONFIG['support_email']}"
                            5.) subject = "Meeting Summary - Pending Review"
                        
                        CASE 2 - If folder_path is "Approved Transcripts" (member scenario):
                            Call send_email with these parameters:
                            1.) text = None
                            2.) file_name = None
                            3.) email_body = "Dear User,<br><br>Thank you for using the Apexon Account Assist for summarizing the meeting notes and actions. Attached is the meeting summary in the Apexon format for your reference.<br><br><a href='{AZURE_CONFIG['approved_folder_url']}' target='_blank'>Link to manage meeting summary</a><br><br>The document has been successfully uploaded to Salesforce.<br><br>Thanks,<br>ServiceAccount"
                            4.) to = USE THE USER EMAIL PROVIDED IN THE WORKFLOW INSTRUCTIONS
                            5.) subject = "Apexon Meeting Summary"
                        
                        After calling send_email, return "TERMINATE" to end the workflow.
                        
                        If you don't receive proper information from Salesforce_uploader agent, don't call send_email and don't return "TERMINATE".
                     """, 
    model_client=oai,
    tools=[send_email_tool],
    description="Sends an email to user or support team based on folder path"
)



#UNINTEGRATED AGENTS
SharepointFileMovementAgent = AssistantAgent(
    "SharepointFileMovementAgent",
    model_client = oai,
    system_message = "You are a sharepoint file mover agent. You have access to move_file_in_sharepoint function." \
    "Your task is to move a particular file from one folder to another folder." \
    "You MUST call the move_file_in_sharepoint tool to move a file from a folder to another folder.",
    description = "Moves a file from one folder to another folder",
    tools = [move_file_in_sharepoint_tool]
)

FileDeletionAgent = AssistantAgent(
    "FileDeletionAgent",
    model_client = oai,
    system_message = "You are a file deletion agent. you have access to delete_file_from_sharepoint function." \
    "Your task is to delete a file from sharepoint. You MUST call the delete_file_from_sharepoint function to delete the file.",
    description="Deletes a particular file from a sharepoint folder.",
    tools = [delete_file_from_sharepoint_tool]
)


#Input Files

#file_name = "meeting_transcript.txt"   
#file_name = "Voicemail_Manjunath.Malagatte@isplahdplayground.onmicrosoft.com_2025-08-18T18_01_22.1347999.txt"  #7         #File with customer name
#file_name = "Customer_name_missing.txt"                                         #File without customer name
#file_name = "Voicemail_Manjunath.Malagatte@isplahdplayground.onmicrosoft.com_2025-08-20T17_28_15.9373900.txt" #11
# file_name = "Adani_Ltd.txt"
#file_name = "Authentication_Changed.txt"
#file_name = "junk_data.txt"
#file_name = "billows.txt"
#file_name = "TFWW.txt"
#file_name = "Account_name.txt"

# transcript = """hi we had a call with LSC Communications customer today.
#  """

transcript = """Today is February 21 2025. So we met with a customer, Anadi Ltd, that is, I repeat, Anadi Ltd. That's the name of the customer we met on 21st of february and the stakeholder name is David watson and james tribbiani. David watson is the principal solution architect. The main priority is What we discussed during this is implementing the new Rendering system for the client and that's the priority as per the stakeholders are concerned. So where the can contribute is we can contribute them in identifying how we can build the Redering System. We can leverage our domain expertise in the logistic field and we can also leverage our technical expertise on the Microsoft text tag and help them to come up with a design of the Rendering system that could be cloud ready. And the next step is setting up a demo with the customer sometime in the next week with our proposed architecture and the solution. Yeah. Overall the meeting went well and client is looking to see our demo and I will upload in the design and the architecture. Thank you."""

# ================== TEAM (GROUP CHAT) ==================


team = SelectorGroupChat(                  
    [
        summarizer_agent,
        GetCustomerListAgent,
        GetCorrectCustomerName_agent,
        extractor_agent,
        FolderPathDecisionAgent,
        UploadToSharepointAgent,
        salesforce_uploader_agent,        
        email_agent,
    ],
    termination_condition=TextMentionTermination("TERMINATE") | MaxMessageTermination(10) | SourceMatchTermination(sources=["Email_agent"]),
    model_client=oai,
    allow_repeated_speaker=False,
)

# ================== WORKFLOW ==================

async def start_r2d2_workflow(user_email: str = "kanchana.k@isplahdplayground.onmicrosoft.com", transcript: str = None):

    workflow_instructions = f"""WORKFLOW: Meeting Transcript Processing

OBJECTIVE: You will be given a meeting transcript text, Process the given meeting transcript through a complete workflow from summarizing the transcript to sending the email.

IMPORTANT: The recipient email address for the final email (for members only) is: {user_email}
For non-members, email will be sent to support team.

REQUIRED WORKFLOW STEPS:
1. START: You are given a meeting transcript text.

2. SUMMARIZE: Have Summarizer agent, first create a summary from the given transcript strictly in this format- 
        1. Date of the meeting
        2. Name of customer
        3. Name & Designation of the customer whom you meet
        4. What are the key priorities for the organization
        5. What are they key priorities of the individual
        6. What are areas of interest where Apexon can contribute
        7. What are the next steps

        After summarizing pass the summary to GetCustomerList agent
 Summarizer agent should not change the name of the customer at any cost on it's own.
 It MUST copy the original name of the customer as it is from the transcript because the task should be completed by CustomerCorrectionAgent.

3.CUSTOMER LIST: Have Get Customer List Agent to get the customer list from sharepoint and then call GetCorrectCustomerNameAgent.
4.Correct CUSTOMER NAME: Have Customer Name Correction Agent which corrects the customer name strictly keeping the original transcript as it is.
 After this agent call Extractor agent.
5. EXTRACT: Have Extractor get key info (CustomerName, Purpose, NextSteps, DateOfMeeting) from the corrected summary. The extractor must return BOTH the corrected summary and the extracted metadata.
    Make sure you also include original transcript content also after summary. It is MUST. After this call FolderPathDecisionAgent.
6. CHECK SITE MEMBERSHIP: Have FolderPathDecisionAgent check if the user ({user_email}) is a member of the SharePoint site.
    - The agent MUST call check_site_membership_and_get_folder function with user_email: {user_email}
    - The function will return True (member) or False (non-member)
    - Based on the result, the AGENT decides the folder_path:
        * If True (member): folder_path = "Approved Transcripts"
        * If False (non-member): folder_path = "Approved Transcripts In Review"
    - The agent must preserve the summary and metadata from Extractor and pass them along with the folder_path AND user_email ({user_email}) to UploadToSharepointAgent.
7. UPLOAD TO SHAREPOINT: Have UploadToSharepoint Agent format and upload the document to SharePoint.
    - UploadToSharepointAgent must call format_summary tool with FOUR arguments:
        * summary (from Extractor)
        * metadata (from Extractor)
        * sub_folder (from FolderPathDecisionAgent - either "Approved Transcripts" or "Approved Transcripts In Review")
        * user_email: {user_email} (to update UserEmail column in SharePoint)
    - The document will be uploaded to the appropriate folder based on user's site membership.
    - The UserEmail column in SharePoint will be set to the logged-in user's email.
    - After upload, pass the response (including folder_path) to Salesforce_uploader.
8. SALESFORCE UPLOAD:
    - Salesforce_uploader will receive file_name, account_name, and folder_path from UploadToSharepointAgent.
    - IF folder_path contains "Review" (non-member): SKIP Salesforce upload and directly notify Email_agent
    - IF folder_path is "Approved Transcripts" (member): Upload to Salesforce, then notify Email_agent
    - After processing, Salesforce_uploader MUST call Email_agent with file_name and folder_path.
    - If salesforce upload fails or if Account name is not found in the salesforce - then stop the workflow strictly
9. EMAIL:
    - Email_agent will check the folder_path to determine recipient:
        * If folder_path contains "Review": Send email to SUPPORT_EMAIL (support team notification)
        * If folder_path is "Approved Transcripts": Send email to {user_email} (user confirmation)
10. Do not call any other agent after Email agent.
11. Make sure to summarize the workflow result in one sentence & return it as the response of the workflow after it ends.

IMPORTANT: Ensure each agent calls their respective function and passes data between steps. The flow is:
Summarizer -> GetCustomerList -> GetCorrectCustomerName -> Extractor -> FolderPathDecisionAgent -> UploadToSharepointAgent -> Salesforce_uploader -> Email_agent

Follow all the steps as mentioned one after another.

Please begin the workflow by directing the Summarizer agent to summarize '{transcript}' from sharepoint site folder.
Make sure that original transcript MUST remain as it is.
Make sure strictly that GetCorectCustomerName agent corrects the customer name by refering to the received sharepoint customer list.
Make sure FolderPathDecisionAgent checks membership for user: {user_email}
Make sure FolderPathDecisionAgent passes user_email ({user_email}) along with summary, metadata, and folder_path to UploadToSharepointAgent.
Make sure UploadToSharepointAgent MUST pass summary as a string, metadata as a dictionary, sub_folder as a string, and user_email as a string ({user_email}) to format_summary tool. 
Make sure you skip the salesforce upload if user is not member of the sharepoint site and document gets uploaded to "Approved Transcripts In Review". In this case it strictly should not get uploaded in salesforce.
Make sure that every agent strictly sticks to their own tasks and do not do any other's agents tasks."""


    result = await Console(team.run_stream(task=workflow_instructions))
    print(result.stop_reason)

  #  print(result.messages[-1].content,"\n")
 

if __name__ == "__main__":
    asyncio.run(start_r2d2_workflow())
